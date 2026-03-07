"""
Reports Page - Generate and download detailed reports
"""
import streamlit as st
from datetime import datetime
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_pdf_report(analysis, include_sections):
    """Generate professional PDF report"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#0369a1'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#0284c7'),
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Title
    story.append(Paragraph("AI Research Intelligence Report", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Document info
    results = analysis['results']
    final_scores = results.get('final_scores', {})
    
    info_data = [
        ['Document:', analysis['filename']],
        ['Analyzed:', analysis['timestamp'].strftime('%Y-%m-%d %H:%M:%S')],
        ['Overall Score:', f"{final_scores.get('overall_score', 0):.0f}/100"]
    ]
    
    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e0f2fe')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#0c4a6e')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#93c5fd'))
    ]))
    story.append(info_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Executive Summary
    if "Executive Summary" in include_sections:
        story.append(Paragraph("Executive Summary", heading_style))
        summary_text = f"""
        This document has been analyzed using our multi-engine plagiarism detection system.
        The overall quality score is {final_scores.get('overall_score', 0):.0f}/100, indicating 
        {'excellent' if final_scores.get('overall_score', 0) >= 80 else 'good' if final_scores.get('overall_score', 0) >= 60 else 'needs improvement'} 
        academic quality.
        """
        story.append(Paragraph(summary_text, styles['BodyText']))
        story.append(Spacer(1, 0.2*inch))
    
    # Scores table
    if "Quality Metrics" in include_sections:
        story.append(Paragraph("Quality Metrics", heading_style))
        
        scores_data = [
            ['Metric', 'Score', 'Status'],
            ['Originality', f"{final_scores.get('originality_score', 0):.0f}/100", 
             'Pass' if final_scores.get('originality_score', 0) >= 70 else 'Needs Work'],
            ['Tone Quality', f"{final_scores.get('tone_score', 0):.0f}/100",
             'Pass' if final_scores.get('tone_score', 0) >= 70 else 'Needs Work'],
            ['Citations', f"{final_scores.get('citation_score', 0):.0f}/100",
             'Pass' if final_scores.get('citation_score', 0) >= 70 else 'Needs Work'],
            ['Structure', f"{final_scores.get('structure_score', 0):.0f}/100",
             'Pass' if final_scores.get('structure_score', 0) >= 70 else 'Needs Work'],
        ]
        
        scores_table = Table(scores_data, colWidths=[2*inch, 1.5*inch, 1.5*inch])
        scores_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0369a1')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f0f9ff')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#93c5fd'))
        ]))
        story.append(scores_table)
        story.append(Spacer(1, 0.3*inch))
    
    # Recommendations
    if "Recommendations" in include_sections and 'guidance' in analysis:
        story.append(Paragraph("Recommendations", heading_style))
        for i, rec in enumerate(analysis['guidance'][:10], 1):
            rec_text = f"{i}. {rec}"
            story.append(Paragraph(rec_text, styles['BodyText']))
            story.append(Spacer(1, 0.1*inch))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_docx_report(analysis, include_sections):
    """Generate professional DOCX report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Research Intelligence Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(3, 105, 161)
    
    # Document info
    doc.add_paragraph()
    results = analysis['results']
    final_scores = results.get('final_scores', {})
    
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_table.rows[0].cells[0].text = 'Document:'
    info_table.rows[0].cells[1].text = analysis['filename']
    info_table.rows[1].cells[0].text = 'Analyzed:'
    info_table.rows[1].cells[1].text = analysis['timestamp'].strftime('%Y-%m-%d %H:%M:%S')
    info_table.rows[2].cells[0].text = 'Overall Score:'
    info_table.rows[2].cells[1].text = f"{final_scores.get('overall_score', 0):.0f}/100"
    
    doc.add_paragraph()
    
    # Executive Summary
    if "Executive Summary" in include_sections:
        doc.add_heading('Executive Summary', 1)
        summary = doc.add_paragraph(
            f"This document has been analyzed using our multi-engine plagiarism detection system. "
            f"The overall quality score is {final_scores.get('overall_score', 0):.0f}/100, indicating "
            f"{'excellent' if final_scores.get('overall_score', 0) >= 80 else 'good' if final_scores.get('overall_score', 0) >= 60 else 'needs improvement'} "
            f"academic quality."
        )
    
    # Quality Metrics
    if "Quality Metrics" in include_sections:
        doc.add_heading('Quality Metrics', 1)
        
        metrics_table = doc.add_table(rows=5, cols=3)
        metrics_table.style = 'Light Grid Accent 1'
        
        # Header
        metrics_table.rows[0].cells[0].text = 'Metric'
        metrics_table.rows[0].cells[1].text = 'Score'
        metrics_table.rows[0].cells[2].text = 'Status'
        
        # Data
        metrics_data = [
            ('Originality', final_scores.get('originality_score', 0)),
            ('Tone Quality', final_scores.get('tone_score', 0)),
            ('Citations', final_scores.get('citation_score', 0)),
            ('Structure', final_scores.get('structure_score', 0))
        ]
        
        for i, (metric, score) in enumerate(metrics_data, 1):
            metrics_table.rows[i].cells[0].text = metric
            metrics_table.rows[i].cells[1].text = f"{score:.0f}/100"
            metrics_table.rows[i].cells[2].text = 'Pass' if score >= 70 else 'Needs Work'
    
    # Recommendations
    if "Recommendations" in include_sections and 'guidance' in analysis:
        doc.add_heading('Recommendations', 1)
        for i, rec in enumerate(analysis['guidance'][:10], 1):
            doc.add_paragraph(f"{i}. {rec}", style='List Number')
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def show():
    """Reports page"""
    
    st.markdown("## 📄 Generate Professional Reports")
    st.markdown("Create detailed PDF or DOCX reports of your analysis")
    
    if 'last_analysis' not in st.session_state or not st.session_state.last_analysis:
        st.info("📭 No analysis available. Analyze a document first!")
        return
    
    analysis = st.session_state.last_analysis
    results = analysis['results']
    final_scores = results.get('final_scores', {})
    
    # Report preview with better styling
    st.markdown("### 📊 Report Preview")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Document", analysis['filename'][:20] + "...")
    with col2:
        st.metric("Overall Score", f"{final_scores.get('overall_score', 0):.0f}/100")
    with col3:
        st.metric("Analyzed", analysis['timestamp'].strftime('%Y-%m-%d'))
    
    st.markdown("---")
    
    # Report options with better UI
    st.markdown("### ⚙️ Report Configuration")
    
    col1, col2 = st.columns(2)
    
    with col1:
        report_format = st.radio(
            "📄 Report Format",
            ["PDF", "DOCX"],
            horizontal=True
        )
        
        st.markdown("**Format Details:**")
        if report_format == "PDF":
            st.info("✓ Professional layout\n✓ Tables and charts\n✓ Print-ready")
        else:
            st.info("✓ Editable format\n✓ Microsoft Word compatible\n✓ Easy to customize")
    
    with col2:
        include_sections = st.multiselect(
            "📑 Include Sections",
            ["Executive Summary", "Quality Metrics", "Recommendations"],
            default=["Executive Summary", "Quality Metrics", "Recommendations"]
        )
        
        st.markdown(f"**Selected:** {len(include_sections)} sections")
    
    st.markdown("---")
    
    # Generate button with better styling
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("🚀 Generate Report", use_container_width=True, type="primary"):
            with st.spinner(f"Generating {report_format} report..."):
                try:
                    if report_format == "PDF":
                        buffer = generate_pdf_report(analysis, include_sections)
                        mime_type = "application/pdf"
                        file_ext = "pdf"
                    else:
                        buffer = generate_docx_report(analysis, include_sections)
                        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        file_ext = "docx"
                    
                    st.success(f"✅ {report_format} report generated successfully!")
                    
                    # Download button
                    st.download_button(
                        label=f"📥 Download {report_format} Report",
                        data=buffer,
                        file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{file_ext}",
                        mime=mime_type,
                        use_container_width=True
                    )
                    
                except Exception as e:
                    st.error(f"❌ Error generating report: {str(e)}")
                    st.info("💡 Make sure reportlab and python-docx are installed:\n```pip install reportlab python-docx```")
